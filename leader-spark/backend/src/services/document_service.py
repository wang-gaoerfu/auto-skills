"""
文档处理服务
"""
import os
import uuid
from pathlib import Path
from typing import List, Dict, Optional, Tuple
from fastapi import UploadFile, HTTPException
from sqlalchemy.ext.asyncio import AsyncSession

from src.core.config import settings
from src.models.document import Document
from src.models.category import Category


class DocumentService:
    """文档处理服务"""

    def __init__(self):
        self.upload_dir = Path(settings.UPLOAD_DIR)
        self.chunk_dir = Path(settings.CHUNK_DIR)
        self.allowed_extensions = set(settings.ALLOWED_FILE_EXTENSIONS)
        self.max_file_size = settings.MAX_FILE_SIZE

        # 确保目录存在
        self.upload_dir.mkdir(parents=True, exist_ok=True)
        self.chunk_dir.mkdir(parents=True, exist_ok=True)

    def validate_file(self, filename: str, file_size: int) -> Tuple[bool, Optional[str]]:
        """
        验证文件

        Args:
            filename: 文件名
            file_size: 文件大小

        Returns:
            (是否有效, 错误消息)
        """
        # 检查文件扩展名
        ext = self.get_file_extension(filename)
        if ext not in self.allowed_extensions:
            return False, f"不支持的文件类型，支持的类型: {', '.join(self.allowed_extensions)}"

        # 检查文件大小
        if file_size > self.max_file_size:
            size_mb = file_size / (1024 * 1024)
            max_mb = self.max_file_size / (1024 * 1024)
            return False, f"文件过大 ({size_mb:.1f}MB)，最大允许 {max_mb:.0f}MB"

        return True, None

    def get_file_extension(self, filename: str) -> str:
        """获取文件扩展名"""
        return Path(filename).suffix.lower().lstrip('.')

    def get_file_type(self, filename: str) -> str:
        """获取文件类型"""
        ext = self.get_file_extension(filename)
        type_mapping = {
            'docx': 'docx',
            'pdf': 'pdf',
            'txt': 'txt',
            'md': 'txt',
        }
        return type_mapping.get(ext, 'unknown')

    async def save_file(self, file: UploadFile, user_id: uuid.UUID) -> Tuple[str, str, int]:
        """
        保存上传的文件

        Args:
            file: 上传的文件
            user_id: 用户 ID

        Returns:
            (存储的文件名, 文件路径, 文件大小)
        """
        # 生成唯一文件名
        ext = self.get_file_extension(file.filename)
        unique_filename = f"{uuid.uuid4()}.{ext}"
        file_path = self.upload_dir / unique_filename

        # 保存文件
        file_size = 0
        with open(file_path, "wb") as f:
            while content := await file.read(1024 * 1024):  # 1MB chunks
                f.write(content)
                file_size += len(content)

        return unique_filename, str(file_path), file_size

    async def extract_text(self, file_path: str, file_type: str) -> str:
        """
        从文件中提取文本

        Args:
            file_path: 文件路径
            file_type: 文件类型

        Returns:
            提取的文本内容
        """
        if file_type == 'docx':
            return await self._extract_from_docx(file_path)
        elif file_type == 'pdf':
            return await self._extract_from_pdf(file_path)
        elif file_type == 'txt':
            return await self._extract_from_txt(file_path)
        else:
            raise ValueError(f"不支持的文件类型: {file_type}")

    async def _extract_from_docx(self, file_path: str) -> str:
        """从 DOCX 文件提取文本"""
        from docx import Document as DocxDocument

        doc = DocxDocument(file_path)
        text_content = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                text_content.append(text)

        # 提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                if row_text.strip():
                    text_content.append(row_text)

        return "\n".join(text_content)

    async def _extract_from_pdf(self, file_path: str) -> str:
        """从 PDF 文件提取文本"""
        import pypdf

        text_content = []
        with open(file_path, 'rb') as file:
            pdf_reader = pypdf.PdfReader(file)
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text and text.strip():
                    text_content.append(text.strip())

        return "\n\n".join(text_content)

    async def _extract_from_txt(self, file_path: str) -> str:
        """从 TXT 文件读取文本"""
        # 尝试多种编码
        encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16']

        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue

        # 如果所有编码都失败，使用错误处理
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()

    def clean_text(self, text: str) -> str:
        """
        清理文本内容

        Args:
            text: 原始文本

        Returns:
            清理后的文本
        """
        import re

        # 移除多余的空白字符
        text = re.sub(r'\s+', ' ', text)

        # 移除特殊字符（保留中文、英文、数字、常用标点）
        text = re.sub(r'[^\u4e00-\u9fff\u3000-\u303fa-zA-Z0-9\s\.,;:!?()（）【】《》、。；：！？]', '', text)

        # 移除过短的行
        lines = text.split('\n')
        lines = [line.strip() for line in lines if len(line.strip()) > 10]

        return '\n'.join(lines)

    def chunk_text(self, text: str, chunk_size: int = None, overlap: int = None) -> List[str]:
        """
        将文本分块

        Args:
            text: 输入文本
            chunk_size: 块大小（字符数）
            overlap: 重叠大小

        Returns:
            文本块列表
        """
        chunk_size = chunk_size or settings.CHUNK_SIZE
        overlap = overlap or settings.CHUNK_OVERLAP

        chunks = []
        start = 0
        text_length = len(text)

        while start < text_length:
            end = start + chunk_size

            # 如果不是最后一块，尝试在句子边界处分割
            if end < text_length:
                # 查找最近的句子结束符
                for delimiter in ['。', '！', '？', '.\n', '!\n', '?\n', '. ', '! ', '? ']:
                    pos = text.rfind(delimiter, start, end)
                    if pos != -1:
                        end = pos + len(delimiter)
                        break

            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)

            # 移动到下一块，考虑重叠
            start = end - overlap if end < text_length else end

        return chunks

    async def delete_file(self, file_path: str) -> bool:
        """
        删除文件

        Args:
            file_path: 文件路径

        Returns:
            是否删除成功
        """
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                return True
            return False
        except Exception as e:
            print(f"删除文件失败: {e}")
            return False

    async def create_document_record(
        self,
        db: AsyncSession,
        filename: str,
        original_filename: str,
        file_path: str,
        file_size: int,
        file_type: str,
        category_id: uuid.UUID,
        user_id: uuid.UUID
    ) -> Document:
        """
        创建文档记录

        Args:
            db: 数据库会话
            filename: 存储的文件名
            original_filename: 原始文件名
            file_path: 文件路径
            file_size: 文件大小
            file_type: 文件类型
            category_id: 分类 ID
            user_id: 上传用户 ID

        Returns:
            文档对象
        """
        document = Document(
            filename=filename,
            original_filename=original_filename,
            file_path=file_path,
            file_size=file_size,
            file_type=file_type,
            category_id=category_id,
            uploaded_by=user_id,
            status="processing"
        )
        db.add(document)
        await db.flush()
        return document


# 创建全局文档服务实例
document_service = DocumentService()
